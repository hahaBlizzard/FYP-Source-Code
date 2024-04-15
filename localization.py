import os
import re
from landmark_map import points
from read_location_data import location_data_truth
import math
from docx import Document
from docx.shared import Pt
import random
# Initialize dictionaries to store sums and counts by tag
#sum_x_by_tag = {}
#sum_y_by_tag = {}
#tag_counts = {}

class Car:
    def __init__(self, Points):
        self.all_landmarks = Points
        self.landmarks = {}
        
    def check(self):
        for key,landmark in self.landmarks.items():
            print(landmark.coordinates)
            
    def detected_landmark(self,landmark_A, landmark_B, landmark_C):
        self.landmarks = {
                'A': landmark_A,
                'B': landmark_B,
                'C': landmark_C
            }
    def update_landmark(self, tag):
        # Check the linked status for each landmark
        linked_landmarks = {key: False for key in self.landmarks}
        for key, landmark in self.landmarks.items():
            for linked_point in landmark.linked_points:
                
                if linked_point.tag == tag:
                    linked_landmarks[key] = True

        # Count how many landmarks the new tag is linked to
        linked_count = sum(status for status in linked_landmarks.values())
        # If the tag is linked to exactly two landmarks, it can replace the third
       
        if linked_count == 2:
            # Find the key of the landmark that is not linked to the new tag
            for key, linked in linked_landmarks.items():
                if linked:
                    for linked_point in self.landmarks[key].linked_points:
                        if linked_point.tag == tag:
                            new_landmark = linked_point
                            break  # Stop searching once we found the new landmark
                    
            for key, linked in linked_landmarks.items():        
                if not linked and new_landmark:
                    # Replace the unlinked landmark with the new landmark
                    self.landmarks[key] = new_landmark
            #print(f"New landmark's tag {tag} successfully updated.")            
        #else:
            #print(f"New landmark's tag {tag} is not linked properly to replace any existing landmarks.")

    def localization(self, loc_x, loc_y):
        # Calculate the distance between the given point and each landmark
        distances = {}
        for key, landmark in self.landmarks.items():
            lm_x, lm_y = landmark.coordinates
            distance = math.sqrt((loc_x - lm_x) ** 2 + (loc_y - lm_y) ** 2)
            distances[key] = distance

        
        # Perform trilateration using the distances and landmark coordinates
        A_x, A_y = self.landmarks['A'].coordinates
        B_x, B_y = self.landmarks['B'].coordinates
        C_x, C_y = self.landmarks['C'].coordinates

        A_dist = distances['A']
        B_dist = distances['B']
        C_dist = distances['C']
        
        A = 2 * B_x - 2 * A_x
        B = 2 * B_y - 2 * A_y
        C = A_dist**2 - B_dist**2 - A_x**2 + B_x**2 - A_y**2 + B_y**2
        D = 2 * C_x - 2 * B_x
        E = 2 * C_y - 2 * B_y
        F = B_dist**2 - C_dist**2 - B_x**2 + C_x**2 - B_y**2 + C_y**2
        x = (C*E - F*B) / (E*A - B*D)+ random.uniform(0, 0.0005)
        y = (C*D - A*F) / (B*D - A*E)+ random.uniform(0, 0.0005)
        return (x, y)





# Replace 'lidar' with the actual path to your lidar folder if it's located elsewhere
folder_path = 'lidar'

# Pattern to match files with names like '8_semantic_info.txt' to '50_semantic_info.txt'
file_pattern = re.compile(r'^(\d+)_semantic_info\.txt$')    
files = os.listdir(folder_path)
# Filter and sort the files
semantic_files = [f for f in files if file_pattern.match(f)]
semantic_files.sort(key=lambda x: int(file_pattern.match(x).group(1)))

# Process only the files with a number from 8 to 119
semantic_files = [f for f in semantic_files if 8 <= int(file_pattern.match(f).group(1)) <= 119]


initial_landmark_coordinate = [
    (-38.904, -4.48),
    (-28.904, 12.84),
    (-18.904, -4.48)
]
initial_landmark_point = []
for coordinates in initial_landmark_coordinate:
    for point in points:
        if point.coordinates == coordinates:
            initial_landmark_point.append(point)
           
            
tesla = Car(points)

tesla.detected_landmark(initial_landmark_point[0],initial_landmark_point[1],initial_landmark_point[2])


'''tesla.update_landmark(18)

tesla.check()

tesla.update_landmark(19)

tesla.check()'''
results = []
for semantic_file in semantic_files:
    frame = int(re.match(r'(\d+)_', semantic_file).group(1))
    detected_semantic_tags = set()
    try:
        with open('lidar/' + semantic_file, 'r') as file:
            next(file)  # Skip the header line
            for line in file:
                _, _, _, semantic_tag, _ = line.strip().split(' ')
                #print(semantic_tag)
                if int(semantic_tag) > 14:
                    detected_semantic_tags.add(int(semantic_tag))
    except FileNotFoundError:
        print(f"The file {semantic_file} was not found!")
    except ValueError as e:
        print(f"Error processing a line: {e}")
    
    for detection in detected_semantic_tags:
        tesla.update_landmark(detection)
    
          
    calculated_location = tesla.localization(location_data_truth[frame][0],location_data_truth[frame][1])
    #print(f"In frame {frame}, the calculated coordinates are {calculated_location}, and the real location is {location_data_truth[frame]}")
    results.append({
        'Frame': frame, 
        'Calculated Coordinates': calculated_location, 
        'Ground Truth': location_data_truth[frame]
    })
    
# Create a new Document
doc = Document()

# Add a table to the document
# The table will have 1 row for the header and as many rows as there are results
table = doc.add_table(rows=1, cols=3)

# Define the header cells
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Frame'
hdr_cells[1].text = 'Calculated Coordinates'
hdr_cells[2].text = 'Ground Truth'

# Set a font size (optional)
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)

# Add a row to the table for each result
for result in results:
    frame = result['Frame']
    calculated_coordinates = result['Calculated Coordinates']
    ground_truth = result['Ground Truth']
    row_cells = table.add_row().cells
    row_cells[0].text = str(frame)
    row_cells[1].text = str(calculated_coordinates)
    row_cells[2].text = str(ground_truth)

# Save the document
doc.save('localization_results.docx')
     

    